from docx import Document
from docx.shared import Inches
import sys


if __name__ == "__main__":
    document = Document()       #creates the document object as well as adds the heading title of the test
    test_name = input("Enter the test name: ")
    document.add_heading(test_name, 0)

    def display():      #function that creates the display for the user where they are prompted for selections
        print("\nTypes of questions:")
        print("\n1. True or False")
        print("\n2. Fill in the Blank")
        print("\n3. Multiple Choice")

        print("\nEnter 'q' to quit")

        selection = input("\nEnter your selection: ")
        selection_list = ['1', '2', '3', 'q', 'Q']      #user selection is then checked against a list of valid options
        
        if selection in selection_list:
            return selection            #returns the user's selection if it is valid and throws an error and restarts the prompt if not
        else:
            print("\nINVALID INPUT - Please enter one of the given options")
            display()

    def true_false():   #function that prompts the user for true/false questions and adds them to the document along with answers
        stop = 0

        while stop == 0:    #while loop that runs until the user inputs !DONE!
                            #the loop prompts the user for the question

            question = input("\nPlease enter the question (e.g. 'The colour of the sky is green') Enter !DONE! if you are done entering questions: ")
            
            if question != "!DONE!":
                p = document.add_paragraph(question, style='List Number')   #adds the question the user inputted as part of a numbered list
                table = document.add_table(rows=1, cols=4)      #creates a table of 4 columns and 1 row with the format of |    |   True    |   |   False   | without the borders
                cell = table.cell(0, 0)
                cell.text = ' '
                cell = table.cell(0, 1)
                cell.text = 'True'
                cell = table.cell(0, 2)
                cell.text = ' '
                cell = table.cell(0, 3)
                cell.text = 'False'
            else:
                stop = 1    #breaks the loop once user inputs '!DONE!

    def fill_in_blank():    #function that prompts the user for fill in the blank questions and adds them to the document
        stop = 0

        while stop == 0:    #while loop works in a similar way to true_false but instead creates just the question as the answer must be written in the blank field
            question = input("\nPlease enter the question (e.g. 'The capital of Ontario is _____________') Enter !DONE! if you are done entering questions: ")

            if question != "!DONE!":
                p = document.add_paragraph(question, style='List Number')
            else:
                stop = 1
            
    def multiple_choice():               #function that prompts the user for multiple choice questions and answers and adds them to the document
        question_stop = 0

        while question_stop == 0:   #while loop that runs until the user no longer wants to input any more questions as indicated by the entry of !DONE!
            question = input("\nPlease enter the question (e.g. '10 x 5 = ?') Enter !DONE! if you are done entering questions: ")
            if question != "!DONE!":
                p = document.add_paragraph(question, style='List Number')   #the question is added as part of the numbered list
                stop = 0
                while stop == 0:    #another while loop that works in the same way but for the answers
                    answer = input("\nEnter an answer choice (enter 'q' if you are done) Enter !DONE! if you are done entering answers: ")
                    if answer == "!DONE!" :
                        stop = 1
                    else:    
                        p = document.add_paragraph(answer, style='List Bullet 2')   #adds the answers under the respective question as bullets with left indentation for clarity
            else:
                question_stop = 1   #breaks the loop once the user inputs !DONE! and ends the function



    program_kill = 0

    while program_kill == 0:    #while loop that will continuously run the program bringing the user back to the main display page until they choose to quit

        user_select = display() #runs the display function and save the user's selection and runs the appropriate function based on that selection

        if user_select == '1':
            true_false()
        elif user_select == '2':
            fill_in_blank()
        elif user_select == '3':
            multiple_choice()
        elif user_select == "q" or user_select == "Q":
            program_kill = 1    #if the user inputs 'q' or 'Q', the while loop breaks

    file_name = test_name.replace(" ", "_") + ".docx"   #names the document file as the test header provided at the beginning of the program replacing any spaces with "_" (e.g. "Test 1" -> "Test_1.docx")

    print("File saved as ", file_name)  #informs user of file name and saves it
    document.save(file_name)